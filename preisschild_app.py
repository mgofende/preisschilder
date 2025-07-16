import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from io import BytesIO
import re

def scrape_product_info(url):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        modell_tag = soup.find('h1', class_='product--title') or soup.find('h1', class_='product-header-title')
        modell = modell_tag.text.strip() if modell_tag else "Modell nicht gefunden"

        artikelnummer = "Artikelnummer nicht gefunden"
        all_text = soup.get_text()
        match = re.search(r'Artikel-?Nr\.?:\s*(\d+)', all_text)
        if match:
            artikelnummer = match.group(1)

        # Erst Standardpreise versuchen
        preis_aktuell_tag = soup.find('span', class_='price--content') or soup.find('div', class_='price--current')
        preis_aktuell = preis_aktuell_tag.text.strip() if preis_aktuell_tag else None

        preis_alt_tag = soup.find('span', class_='price--line-through') or soup.find('span', class_='price-old')
        preis_alt = preis_alt_tag.text.strip() if preis_alt_tag else None

        # Fallback: Meta-Tag mit itemprop="price"
        if not preis_aktuell:
            meta_price = soup.find('meta', itemprop='price')
            if meta_price and meta_price.has_attr('content'):
                preis_aktuell = meta_price['content'].strip() + " €"

        # Sternchen entfernen
        if preis_aktuell:
            preis_aktuell = preis_aktuell.replace('*', '').strip()
        if preis_alt:
            preis_alt = preis_alt.replace('*', '').strip()

        # Bild-URL finden
        img_url = None
        img_tag = soup.find('img', attrs={'data-img-large': True})
        if img_tag:
            img_url = img_tag['data-img-large']
        else:
            match_img = re.findall(r'data-img-large="(https://[^"]+\.jpg)"', response.text)
            if match_img:
                img_url = match_img[0]

        return modell, artikelnummer, preis_aktuell, preis_alt, img_url

    except Exception as e:
        st.error(f"Fehler beim Auslesen der Webseite: {e}")
        return None, None, None, None, None

def create_word_file(modell, artikelnummer, preis_aktuell, preis_alt, img_url):
    doc = Document()
    section = doc.sections[0]

    # A4 Hochformat
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = WD_ORIENT.PORTRAIT

    # Ränder setzen
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(31)
    section.right_margin = Mm(31)

    # Schriftart Arial global
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    rPr = style.element.rPr
    rFonts = rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Arial')

    # Hintergrundbild (A5) mittig - NEU mit anderem Link und Fallback
    bg_url = "https://backend.ofen.de/media/image/63/2e/5c/Grafik-fuer-Preisschildchen-unten.png"
    try:
        bg_response = requests.get(bg_url)
        bg_response.raise_for_status()
        bg_stream = BytesIO(bg_response.content)

        p_bg = doc.add_paragraph()
        p_bg.alignment = 1  # zentriert
        run_bg = p_bg.add_run()
        run_bg.add_picture(bg_stream, width=Mm(148), height=Mm(210))  # A5 Größe
    except:
        # Fallback URL
        bg_url_fallback = "https://www.ofen.de/media/image/63/2e/5c/Grafik-fuer-Preisschildchen-unten.png"
        try:
            bg_response = requests.get(bg_url_fallback)
            bg_response.raise_for_status()
            bg_stream = BytesIO(bg_response.content)

            p_bg = doc.add_paragraph()
            p_bg.alignment = 1
            run_bg = p_bg.add_run()
            run_bg.add_picture(bg_stream, width=Mm(148), height=Mm(210))
        except:
            doc.add_paragraph("Hintergrundbild konnte nicht geladen werden.")

    # Produktbild zentriert und 80 mm breit
    if img_url:
        try:
            img_response = requests.get(img_url)
            img_response.raise_for_status()
            img_stream = BytesIO(img_response.content)

            p_img = doc.add_paragraph()
            p_img.alignment = 1
            run_img = p_img.add_run()
            run_img.add_picture(img_stream, width=Mm(80))
        except:
            doc.add_paragraph("Produktbild konnte nicht geladen werden.")
    else:
        doc.add_paragraph("Kein Produktbild verfügbar.")

    # Text zentriert
    p = doc.add_paragraph()
    p.alignment = 1

    run1 = p.add_run(modell + "\n")
    run1.font.size = Pt(18)
    run1.font.bold = True

    run2 = p.add_run(f"Artikelnummer: {artikelnummer}\n")
    run2.font.size = Pt(11)

    # Kleiner Abstand zwischen Artikelnummer und Preis
    run_spacer = p.add_run(" \n")
    run_spacer.font.size = Pt(4)

    run3 = p.add_run(preis_aktuell + "\n")
    run3.font.size = Pt(24)
    run3.font.bold = True
    run3.font.color.rgb = RGBColor(200, 0, 0)

    if preis_alt:
        run4 = p.add_run(preis_alt)
        run4.font.size = Pt(16)
        run4.font.strike = True
        run4.font.color.rgb = RGBColor(120, 120, 120)

    word_io = BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io

# Streamlit UI
st.set_page_config(page_title="Preisschild Generator A5 auf A4 mit neuem Hintergrund", page_icon="🧾")
st.title("🧾 Preisschild Generator mit neuem Backend-Hintergrund (A5 auf A4)")

st.markdown("**Gib den Produktlink von Ofen.de ein:**")

url = st.text_input("🔗 Produkt-URL eingeben:")

if url:
    modell, artikelnummer, preis_aktuell, preis_alt, img_url = scrape_product_info(url)

    if modell and artikelnummer and preis_aktuell:
        st.success("✅ Produktdaten erfolgreich geladen!")
        st.markdown(f"**Modell:** {modell}")
        st.markdown(f"**Artikelnummer:** {artikelnummer}")
        st.markdown(f"**Preis:** {preis_aktuell}")
        if preis_alt:
            st.markdown(f"**Alter Preis:** ~~{preis_alt}~~")
        if img_url:
            st.image(img_url, width=300)

        if st.button("📄 Preisschild erstellen"):
            file = create_word_file(modell, artikelnummer, preis_aktuell, preis_alt, img_url)
            if file:
                st.download_button(
                    label="⬇️ Preisschild als Word herunterladen",
                    data=file,
                    file_name="preisschild_A5_auf_A4_neuer_Hintergrund.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.error("❌ Einige Produktdaten konnten nicht geladen werden.")